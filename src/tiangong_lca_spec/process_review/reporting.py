"""Utilities for rendering review reports into structured documents."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Mapping, Sequence

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .models import ProcessReviewResult, ReviewFinding


def generate_docx_report(
    result: ProcessReviewResult,
    *,
    dataset_metadata: Mapping[str, str | None],
    context: Mapping[str, object],
    output_path: Path,
    template_path: str | Path | None = None,
) -> None:
    """Render the review findings into a DOCX report aligned with the official template."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if template_path:
        template = Path(template_path)
        if template.exists():
            document = Document(template)
        else:
            document = Document()
    else:
        document = Document()

    _configure_document_layout(document)
    _configure_document_styles(document)

    if not document.paragraphs:
        _initialise_document(document)

    _write_general_information_section(document, result, dataset_metadata, context)
    _write_compliance_assessment_section(document, context, result)
    _write_findings_section(document, result)
    _write_references_section(document, context)
    _write_declaration_section(document, context, result)

    document.save(str(output_path))


def _initialise_document(document: Document) -> None:
    _add_heading(document, "审查报告", level=0)
    document.add_paragraph("REVIEW REPORTING")


def _write_general_information_section(
    document: Document,
    result: ProcessReviewResult,
    dataset_metadata: Mapping[str, str | None],
    context: Mapping[str, object],
) -> None:
    general = dict(context.get("general_information") or {})

    dataset_name = (
        general.get("dataset_name")
        or dataset_metadata.get("name")
        or dataset_metadata.get("uuid")
        or "待补充"
    )
    dataset_uuid = (
        general.get("dataset_uuid")
        or dataset_metadata.get("uuid")
        or "待补充"
    )
    dataset_version = general.get("dataset_version") or dataset_metadata.get("version")
    uuid_text = dataset_uuid if not dataset_version else f"{dataset_uuid} / {dataset_version}"

    dataset_locator = general.get("dataset_locator") or dataset_metadata.get("locator") or "待补充"
    commissioners = general.get("commissioners") or []
    if isinstance(commissioners, str):
        commissioners = [commissioners]
    commissioner_text = "\n".join(str(item) for item in commissioners) or "待补充"

    reviewers = _format_reviewers(general.get("reviewers"))
    reviewer_type = general.get("reviewer_type") or result.metadata.review_type

    review_type_detail = general.get("review_type_detail") or result.metadata.review_type
    method_scope = general.get("method_scope") or f"{result.metadata.method_label} / {result.metadata.scope}"
    completion_date = general.get("review_completion_date") or datetime.utcnow().date().isoformat()
    compliance_system = general.get("compliance_system") or "待补充"
    ef_reference = general.get("ef_reference_package") or "待补充"

    _add_heading(document, "一般信息 / General information", level=1)
    table = document.add_table(rows=1, cols=2)
    _set_table_style(table, "Table Grid")
    header_cells = table.rows[0].cells
    header_cells[0].text = "项目 / Item"
    header_cells[1].text = "内容 / Details"
    _style_table_header_row(table.rows[0])

    entries = [
        ("数据集名称 / Data set name", dataset_name),
        ("数据集 UUID 和版本号 / Data set UUID and version number", uuid_text),
        ("数据集位置 (URI/联系点) / Data set locator", dataset_locator),
        ("审查委托人 / Review commissioner(s)", commissioner_text),
        ("审查员（含单位、联系方式） / Reviewer(s) & affiliation", reviewers),
        ("审查员类型 / Reviewer type", reviewer_type),
        ("适用的审查类型 / Review type applied", review_type_detail),
        ("审查方法及范围 / Method used for review & scope", method_scope),
        ("审查完成日期 / Date of review completion", str(completion_date)),
        ("审查依据 / Compliance system name", compliance_system),
        ("EF 参考包兼容性 / EF reference package compatibility", ef_reference),
    ]

    for label, value in entries:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value)
        _apply_cell_style(row_cells[0], "ReviewTableLabel")
        _apply_cell_style(row_cells[1], "ReviewTableValue")


def _write_compliance_assessment_section(
    document: Document,
    context: Mapping[str, object],
    result: ProcessReviewResult,
) -> None:
    categories = [
        ("ef_requirements", "与特定 EF 要求的一致性 / Compliance with specific EF requirements"),
        ("allocation_rules", "分配规则的明确性和一致性 / Allocation rules clarity"),
        ("cff", "循环足迹公式（正确实施） / Circular Footprint Formula"),
        ("lcia_results", "LCIA 结果一致性 / LCIA results consistency"),
        ("nomenclature", "命名法 / Nomenclature"),
        ("documentation", "文件记录 / Documentation"),
        ("ilcd_format", "文件格式的适当性（ILCD 格式） / ILCD documentation format"),
        ("ilcd_validator", "通过 ILCD 验证器的验证 / Validation by ILCD validator"),
        ("data_quality", "数据质量标准和评级 / Data quality rating"),
        ("cut_off", "截断法（Cut-off）"),
        ("additional_info", "补充信息 / Additional information"),
    ]
    compliance_ctx = dict(context.get("compliance_assessment") or {})

    _add_heading(document, "总体合规评估 / Overall compliance assessment", level=1)
    table = document.add_table(rows=1, cols=3)
    _set_table_style(table, "Table Grid")
    header = table.rows[0].cells
    header[0].text = "内容 / Aspect"
    header[1].text = "是否 / Status"
    header[2].text = "备注 / Comments"
    _style_table_header_row(table.rows[0])

    default_status = "Yes" if not result.review_findings else "Pending"

    for key, label in categories:
        entry = compliance_ctx.get(key) or {}
        if isinstance(entry, str):
            status = entry
            comments = ""
        else:
            status = entry.get("status", default_status)
            comments = entry.get("comments", "")
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(status)
        row[2].text = comments or "-"
        _apply_cell_style(row[0], "ReviewTableLabel")
        _apply_cell_style(row[1], "ReviewTableStatus")
        _apply_cell_style(row[2], "ReviewTableValue")


def _write_findings_section(document: Document, result: ProcessReviewResult) -> None:
    _add_heading(document, "发现 / Findings", level=1)
    findings: list[ReviewFinding] = result.validation_findings + result.review_findings
    if not findings:
        document.add_paragraph("未记录自动化审查发现。")
        return

    for finding in findings:
        text = f"[{finding.severity.upper()}] {finding.category}: {finding.message}"
        if finding.path:
            text += f" (path: {finding.path})"
        paragraph = _add_paragraph_with_style(document, text, style="ReviewBullet", bullet=True)
        if finding.evidence:
            paragraph.add_run(f"\n证据 / Evidence: {finding.evidence}")
        if finding.suggestion:
            paragraph.add_run(f"\n建议 / Suggestion: {finding.suggestion}")


def _write_references_section(document: Document, context: Mapping[str, object]) -> None:
    references = context.get("references") or []
    if isinstance(references, str):
        references = [references]

    _add_heading(document, "审查参考资料 / Documents consulted", level=1)
    if not references:
        document.add_paragraph("待补充。")
    else:
        for item in references:
            _add_paragraph_with_style(document, str(item), style="ReviewBullet", bullet=True)


def _write_declaration_section(
    document: Document,
    context: Mapping[str, object],
    result: ProcessReviewResult,
) -> None:
    declaration = context.get("declaration")
    if not declaration:
        compliant = not any(f.severity == "error" for f in result.review_findings)
        declaration = (
            "审查人声明：根据当前自动化审查结果，该数据集初步判定为符合环境足迹规则。最终结论有待人工复核确认。"
            if compliant
            else "审查人声明：根据当前自动化审查结果，该数据集存在需整改的问题，未满足环境足迹合规要求。"
        )

    _add_heading(document, "合规声明 / Compliance declaration", level=1)
    document.add_paragraph(str(declaration))

    signatures = context.get("signature") or {}
    reviewers = signatures.get("reviewers") if isinstance(signatures, dict) else signatures
    if reviewers:
        _add_heading(document, "签字 / Signatures", level=1)
        if isinstance(reviewers, str):
            reviewers = [reviewers]
        for entry in reviewers:
            document.add_paragraph(str(entry))


def _format_reviewers(raw: object) -> str:
    if not raw:
        return "待补充"
    if isinstance(raw, str):
        return raw
    if isinstance(raw, Sequence):
        rendered: list[str] = []
        for item in raw:
            if isinstance(item, Mapping):
                pieces = [
                    str(item.get("name") or ""),
                    str(item.get("affiliation") or ""),
                    str(item.get("contact") or ""),
                    str(item.get("reviewer_type") or ""),
                ]
                rendered.append(" / ".join(filter(None, pieces)))
            else:
                rendered.append(str(item))
        return "\n".join(rendered) or "待补充"
    if isinstance(raw, Mapping):
        return " / ".join(str(raw.get(part) or "") for part in ("name", "affiliation", "contact"))
    return str(raw)


def _add_heading(document: Document, text: str, level: int) -> None:
    try:
        document.add_heading(text, level=level)
    except KeyError:
        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in paragraph.runs:
            run.bold = True


def _set_table_style(table, style_name: str | None) -> None:
    candidates: list[str | None] = ["ReviewTable"]
    if style_name and style_name not in candidates:
        candidates.append(style_name)
    candidates.extend(["Table Grid", "Normal Table", None])

    for candidate in candidates:
        if candidate is None:
            continue
        try:
            table.style = candidate
            break
        except (KeyError, ValueError, AttributeError):
            continue

    table.autofit = True
    try:
        table.allow_autofit = True
    except AttributeError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def _style_table_header_row(row) -> None:
    for cell in row.cells:
        _apply_cell_style(cell, "ReviewTableHeader")


def _apply_cell_style(cell, style_name: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph("")
    for paragraph in cell.paragraphs:
        paragraph.style = style_name


def _add_paragraph_with_style(document: Document, text: str, style: str, *, bullet: bool = False):
    try:
        return document.add_paragraph(text, style=style)
    except KeyError:
        prefix = "• " if bullet else ""
        return document.add_paragraph(f"{prefix}{text}")


def _configure_document_layout(document: Document) -> None:
    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)
        section.different_first_page_header_footer = False
        section.gutter = Cm(0)


def _configure_document_styles(document: Document) -> None:
    styles = document.styles

    def apply_font(style, size: int, *, bold: bool = False) -> None:
        style.font.name = "宋体"
        style.font.size = Pt(size)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    normal = styles["Normal"]
    apply_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.2

    title = _ensure_paragraph_style(styles, "Title")
    apply_font(title, 20, bold=True)
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    title.paragraph_format.keep_with_next = True

    for heading_name, size in (("Heading 1", 16), ("Heading 2", 14)):
        heading = _ensure_paragraph_style(styles, heading_name)
        apply_font(heading, size, bold=True)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.keep_together = True

    bullet_style = _ensure_paragraph_style(styles, "ReviewBullet")
    try:
        bullet_style.base_style = styles["List Bullet"]
    except KeyError:
        bullet_style.base_style = styles["Normal"]
    apply_font(bullet_style, 11)
    bullet_style.paragraph_format.left_indent = Pt(18)
    bullet_style.paragraph_format.first_line_indent = Pt(0)
    bullet_style.paragraph_format.space_before = Pt(0)
    bullet_style.paragraph_format.space_after = Pt(4)
    bullet_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    table_style = _ensure_table_style(styles, "ReviewTable")
    table_style.font.name = "宋体"
    table_style.font.size = Pt(11)
    try:
        tbl_pr = table_style.element.get_or_add_tblPr()
    except AttributeError:
        tbl_pr = None
    if tbl_pr is not None:
        tbl_borders = tbl_pr.find(qn("w:tblBorders"))
        if tbl_borders is None:
            tbl_borders = OxmlElement("w:tblBorders")
            tbl_pr.append(tbl_borders)
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = tbl_borders.find(qn(f"w:{border_name}"))
            if border is None:
                border = OxmlElement(f"w:{border_name}")
                tbl_borders.append(border)
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "8")
            border.set(qn("w:color"), "000000")

    table_header = _ensure_paragraph_style(styles, "ReviewTableHeader")
    table_header.base_style = styles["Normal"]
    apply_font(table_header, 11, bold=True)
    header_format = table_header.paragraph_format
    header_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_format.space_before = Pt(0)
    header_format.space_after = Pt(4)
    header_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    header_format.keep_together = True

    table_label = _ensure_paragraph_style(styles, "ReviewTableLabel")
    table_label.base_style = styles["Normal"]
    apply_font(table_label, 11, bold=True)
    label_format = table_label.paragraph_format
    label_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    label_format.space_before = Pt(0)
    label_format.space_after = Pt(4)
    label_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    label_format.keep_together = True

    table_value = _ensure_paragraph_style(styles, "ReviewTableValue")
    table_value.base_style = styles["Normal"]
    apply_font(table_value, 11)
    value_format = table_value.paragraph_format
    value_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    value_format.space_before = Pt(0)
    value_format.space_after = Pt(4)
    value_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    value_format.keep_together = True

    table_status = _ensure_paragraph_style(styles, "ReviewTableStatus")
    table_status.base_style = styles["Normal"]
    apply_font(table_status, 11)
    status_format = table_status.paragraph_format
    status_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    status_format.space_before = Pt(0)
    status_format.space_after = Pt(4)
    status_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    status_format.keep_together = True


def _ensure_paragraph_style(styles, name: str):
    try:
        return styles[name]
    except KeyError:
        return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _ensure_table_style(styles, name: str):
    try:
        return styles[name]
    except KeyError:
        return styles.add_style(name, WD_STYLE_TYPE.TABLE)
